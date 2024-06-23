from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from htmldocx import HtmlToDocx


def docx_generate(self, outline, objectives, outcome_standards, evaluation_tests):
    doc = Document()
    parser = HtmlToDocx()
    styles = doc.styles["Normal"]
    font = styles.font
    font.name = "Times New Roman"
    font.size = Pt(13)

    p1 = doc.add_paragraph()
    p1.add_run("TRƯỜNG ĐẠI HỌC MỞ THÀNH PHỐ HỒ CHÍ MINH").bold = True
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p2 = doc.add_paragraph()
    p2.add_run("KHOA CÔNG NGHỆ THÔNG TIN").bold = True
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p3 = doc.add_paragraph()
    p3.add_run("ĐỀ CƯƠNG MÔN HỌC").bold = True
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p4 = doc.add_paragraph(style="List Number")
    p4.add_run("Thông tin tổng quát").bold = True

    p5 = doc.add_paragraph(style="List Bullet 2")
    p5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p5.add_run("Tên môn học tiếng Việt: ").bold = True
    p5.add_run(outline.get("subject").get("vie_name").upper())

    p6 = doc.add_paragraph(style="List Bullet 2")
    p6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY_LOW
    p6.add_run("Tên môn học tiếng Anh: ").bold = True
    p6.add_run(outline.get("subject").get("eng_name").upper())

    p7 = doc.add_paragraph(style="List Bullet 2")
    p7.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p7.add_run("Thuộc khối kiến thức/kỹ năng: ").bold = True
    p7.add_run(outline.get("subject").get("knowledge_base"))

    p8 = doc.add_paragraph(style="List Bullet 2")
    p8.add_run("Số tín chỉ").bold = True

    table1 = doc.add_table(rows=2, cols=4)
    table1.alignment = WD_TABLE_ALIGNMENT.CENTER
    table1.style = "Table Grid"
    cell = table1.rows[0].cells
    cell[0].text = "Tổng số"
    cell[1].text = "Lý thuyết"
    cell[2].text = "Thực hành"
    cell[3].text = "Tự học"
    cell_1 = table1.rows[1].cells
    cell_1[0].text = str(outline.get("subject").get("theory_credit") + outline.get("subject").get(
        "practical_credit"))
    cell_1[1].text = str(outline.get("subject").get("theory_credit"))
    cell_1[2].text = str(outline.get("subject").get("practical_credit"))
    cell_1[3].text = "{}"

    p9 = doc.add_paragraph(style="List Bullet 2")
    p9.add_run("Phụ trách môn học").bold = True

    p10 = doc.add_paragraph(style="List Bullet 3")
    p10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p10.add_run("Phụ trách: {}")

    p11 = doc.add_paragraph(style="List Bullet 3")
    p11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p11.add_run("Giảng viên: {}")

    p12 = doc.add_paragraph(style="List Bullet 3")
    p12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p12.add_run("Email: {}")

    p13 = doc.add_paragraph(style="List Number")
    p13.add_run("Thông tin môn học").bold = True

    p14 = doc.add_paragraph(style="List Bullet 2")
    p14.add_run("Mô tả môn học").bold = True

    description = outline.get("description")
    parser.add_html_to_document(description, doc)

    p16 = doc.add_paragraph(style="List Bullet 2")
    p16.add_run("Môn học điều kiện").bold = True

    table2 = doc.add_table(rows=1, cols=3)
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    table2.style = "Table Grid"
    cell = table2.rows[0].cells
    cell[0].text = "STT"
    cell[1].text = "Môn học điều kiện"
    cell[2].text = "Mã môn học"

    insistences = outline.get("subject").get("insistence")
    count = 1
    for insis in insistences:
        cell = table2.add_row().cells
        cell[0].text = str(count)
        cell[1].text = insis.get("vie_name")
        cell[2].text = insis.get("subject_id")
        count += 1

    p17 = doc.add_paragraph(style="List Bullet 2")
    p17.add_run("Mục tiêu môn học").bold = True

    p18 = doc.add_paragraph()
    p18.add_run(
        "Môn học cung cấp cho người học những kiến thức, kỹ năng cũng như cho người học có các thái độ như sau:")

    table3 = doc.add_table(rows=1, cols=3)
    table3.alignment = WD_TABLE_ALIGNMENT.CENTER
    table3.style = "Table Grid"
    cell = table3.rows[0].cells
    cell[0].text = "Mục tiêu môn học"
    cell[1].text = "Mô tả"
    cell[2].text = "CĐR CTĐT phân bổ cho môn học"
    for objective in objectives:
        cell = table3.add_row().cells
        cell[0].text = objective.get("name")
        parser.add_html_to_cell(objective.get("description"), cell[1])
        parser.add_html_to_cell(objective.get("standard"), cell[2])

    p19 = doc.add_paragraph(style="List Bullet 2")
    p19.add_run("Chuẩn đầu ra(CĐR) môn học").bold = True

    p20 = doc.add_paragraph()
    p20.add_run("Học xong môn học này, sinh viên sẽ đạt các chuẩn đầu ra sau:")

    table4 = doc.add_table(rows=1, cols=3)
    table4.alignment = WD_TABLE_ALIGNMENT.CENTER
    table4.style = "Table Grid"
    cell = table4.rows[0].cells
    cell[0].text = "Mục tiêu môn học"
    cell[1].text = "CĐR môn học(CLO)"
    cell[2].text = "Mô tả CĐR"
    for standard in outcome_standards:
        cell = table4.add_row().cells
        cell[0].text = standard.get("subject_objective")
        cell[1].text = standard.get("name")
        parser.add_html_to_cell(standard.get("description"), cell[2])

    p21 = doc.add_paragraph(style="List Bullet 2")
    p21.add_run("Học liệu").bold = True

    document = outline.get("document")
    parser.add_html_to_document(document, doc)

    p23 = doc.add_paragraph(style="List Bullet 2")
    p23.add_run("Phương pháp giảng dạy - học tập").bold = True

    teaching_method = outline.get("teaching_method")
    parser.add_html_to_document(teaching_method, doc)

    p24 = doc.add_paragraph(style="List Bullet 2")
    p24.add_run("Đánh giá môn học").bold = True

    table5 = doc.add_table(rows=2, cols=5)
    table5.style = "Table Grid"
    cell = table5.rows[0].cells
    cell_1 = table5.rows[1].cells
    cell[0].text = "Thành phần đánh giá/Type of accessment"
    cell[1].text = "Bài đánh giá/Accessment methods"
    cell[2].text = "Thời điểm/Accessment time"
    cell[3].text = "CĐR môn học/CLOs"
    cell[4].text = "Tỷ lệ %/Weight %"
    cell_1[0].text = "(1)"
    cell_1[1].text = "(2)"
    cell_1[2].text = "(3)"
    cell_1[3].text = "(4)"
    for test in evaluation_tests:
        cell = table5.add_row().cells
        cell[0].text = test.get("test_type")
        cell[1].text = test.get("name")
        parser.add_html_to_cell(test.get("time"), cell[2])
        cell[3].text = ""
        for i in test.get("subject_standard"):
            cell[3].text = cell[3].text + " " + i.get("name")
        cell[4].text = str(test.get("score_ratio"))

    ### canh giua
    for table in doc.tables:
        for col in table.columns:
            for cell in col.cells:
                for par in cell.paragraphs:
                    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ###
    return doc
